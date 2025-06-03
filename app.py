import streamlit as st
import pandas as pd
import random
import datetime
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

st.title("Event Feedback Generator")

num_students = st.number_input("Number of students", min_value=1, step=1)
event_name = st.text_input("Event name")
event_date_str = st.text_input("Date of event (DD-MM-YYYY)")

# Initialize session_state keys if not present
if "files_generated" not in st.session_state:
    st.session_state.files_generated = False
    st.session_state.excel_buffer = None
    st.session_state.word_buffer = None

if st.button("Generate Feedback Files"):
    try:
        event_date = datetime.datetime.strptime(event_date_str, "%d-%m-%Y")
    except:
        st.error("Invalid date format, use DD-MM-YYYY")
        st.stop()

    questions = [
        "1. How satisfied were you with the overall event?",
        "2. How well was the event organized?",
        "3. How informative did you find the sessions?",
        "4. How would you rate the event venue and facilities?",
        "5. How likely are you to recommend this event to others?"
    ]

    student_names = [f"Student {i+1}" for i in range(num_students)]
    student_emails = [f"student{i+1}@example.com" for i in range(num_students)]
    student_ids = [f"BWU{i+1:04d}" for i in range(num_students)]

    timestamps = []
    used_datetimes = set()
    while len(timestamps) < num_students:
        days_after = random.randint(1, 10)
        random_hour = random.randint(10, 17)
        random_minute = random.randint(0, 59)
        random_second = random.randint(0, 59)
        dt = event_date + datetime.timedelta(days=days_after, hours=random_hour, minutes=random_minute, seconds=random_second)
        if dt not in used_datetimes:
            used_datetimes.add(dt)
            timestamps.append(dt)
    timestamps.sort()

    data = []
    for i in range(num_students):
        timestamp = timestamps[i]
        feedback = [random.randint(2, 5) for _ in questions]
        row = [timestamp.strftime("%Y-%m-%d %H:%M:%S"), student_emails[i], student_names[i], student_ids[i]] + feedback
        data.append(row)

    columns = ["Timestamp", "Email", "Name", "BWU Student Code"] + questions
    df = pd.DataFrame(data, columns=columns)

    # Prepare Excel in memory
    excel_buffer = BytesIO()
    df.to_excel(excel_buffer, index=False)
    excel_buffer.seek(0)  # Reset pointer

    # Prepare Word document in memory
    doc = Document()
    doc.add_heading(f'{event_name} - Event Feedback Summary', 0)
    doc.add_paragraph(f"Date of Event: {event_date_str}")
    doc.add_paragraph(f"Number of Participants: {num_students}")

    colors = ['red', 'green', 'blue', 'orange', 'purple']

    for idx, question in enumerate(questions):
        ratings = df[question].tolist()
        plt.figure(figsize=(6, 4))
        plt.hist(ratings, bins=[0.5,1.5,2.5,3.5,4.5,5.5], align='mid', rwidth=0.8,
                 color=colors[idx % len(colors)], edgecolor='black')
        plt.xlabel("Rating (Likert Scale: 1 to 5)")
        plt.ylabel("Number of Responses")
        plt.title(question)
        plt.xticks([1, 2, 3, 4, 5])
        plt.grid(True, linestyle='--', alpha=0.6)
        for spine in plt.gca().spines.values():
            spine.set_linewidth(2)

        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', bbox_inches='tight')
        plt.close()
        img_buffer.seek(0)

        doc.add_heading(question, level=2)
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(img_buffer, width=Inches(5.5))
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)  # Reset pointer

    # Save generated files in session_state to keep buttons visible after generation
    st.session_state.files_generated = True
    st.session_state.excel_buffer = excel_buffer
    st.session_state.word_buffer = word_buffer

    st.success("Files generated successfully!")

# Show download buttons if files are generated
if st.session_state.files_generated:
    st.download_button("Download Excel", data=st.session_state.excel_buffer,
                       file_name=f"{event_name}_feedback.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.download_button("Download Word Document", data=st.session_state.word_buffer,
                       file_name=f"{event_name}_feedback.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
