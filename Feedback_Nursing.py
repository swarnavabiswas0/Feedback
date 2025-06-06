import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
import os

st.set_page_config(layout="wide")
st.title("📊 Event Feedback Analysis Report Generator")

# Full questions to appear in Word
QUESTION_TEXTS = {
    "1. Overall Rating": "Overall Rating: Please rate the overall quality of the event on a scale of 1 to 5 (1 being Poor, 5 being Excellent)",
    "2. Objectives Met": "Event/Activity Objectives: Were the objectives of the event clearly communicated and met?",
    "3. Event Organization": "How well was the event/activity organized?",
    "4. Interaction and Engagement": "Interaction and Engagement",
    "5. Logistics": "How would you rate the Logistics?"
}

HEADERS = ["Overall Rating", "Objectives", "Organization", "Interaction", "Logistics", "Comments"]

# Upload file
uploaded_file = st.file_uploader("Upload Excel or CSV file", type=["xlsx", "csv"])
if uploaded_file:
    # Read file
    if uploaded_file.name.endswith('.csv'):
        df = pd.read_csv(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file)

    # Auto-rename columns
    def clean_columns(col):
        col = col.lower()
        if "overall" in col:
            return "Overall Rating"
        elif "objective" in col:
            return "Objectives"
        elif "organize" in col:
            return "Organization"
        elif "interaction" in col:
            return "Interaction"
        elif "logistics" in col:
            return "Logistics"
        elif "comment" in col:
            return "Comments"
        return col

    df.columns = [clean_columns(c) for c in df.columns]

    # Show preview
    st.subheader("🔍 Data Preview")
    st.dataframe(df.head())

    # Convert ratings
    df['Overall Rating'] = pd.to_numeric(df['Overall Rating'], errors='coerce').fillna(0).astype(int)

    # Plot generator
    def plot_chart(data, title, kind='bar'):
        fig, ax = plt.subplots()
        if kind == 'pie':
            data.value_counts().plot(kind='pie', autopct='%1.1f%%', startangle=90, ax=ax)
            ax.set_ylabel('')
        else:
            data.value_counts().sort_index().plot(kind='bar', color='skyblue', edgecolor='black', ax=ax)
            ax.set_xlabel(title)
            ax.set_ylabel("Responses")
        ax.set_title(title)
        st.pyplot(fig)
        return fig

    # Generate and show charts
    st.subheader("📈 Feedback Charts")

    fig1 = plot_chart(df['Overall Rating'], "1. Overall Rating")
    fig2 = plot_chart(df['Objectives'], "2. Objectives Met", kind='pie')
    fig3 = plot_chart(df['Organization'], "3. Event Organization")
    fig4 = plot_chart(df['Interaction'], "4. Interaction and Engagement")
    fig5 = plot_chart(df['Logistics'], "5. Logistics")

    # Word Report Generation
    if st.button("📥 Generate Word Report"):
        doc = Document()
        doc.add_heading("Event Feedback Analysis Report", 0)

        def add_section(title, fig):
            doc.add_heading(title, level=1)
            doc.add_paragraph(QUESTION_TEXTS[title])
            image_stream = BytesIO()
            fig.savefig(image_stream, format='png')
            image_stream.seek(0)
            doc.add_picture(image_stream, width=Inches(5.5))
            image_stream.close()

        add_section("1. Overall Rating", fig1)
        add_section("2. Objectives Met", fig2)
        add_section("3. Event Organization", fig3)
        add_section("4. Interaction and Engagement", fig4)
        add_section("5. Logistics", fig5)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Report generated successfully.")
        st.download_button(label="📄 Download Word Report",
                           data=buffer,
                           file_name="Feedback_Analysis_Report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
